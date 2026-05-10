import csv
import textwrap
import urllib.parse
import webbrowser
from datetime import datetime
from pathlib import Path

import fitz


def _wrap(text: str, width: int = 90) -> list[str]:
    return textwrap.wrap(text, width) or [""]


class ExportService:

    def export_pdf(self, path: str, diagnosis: dict, hypotheses: list) -> None:
        doc = fitz.open()
        page = doc.new_page(width=595, height=842)

        black = (0, 0, 0)
        gray = (0.35, 0.35, 0.35)
        light = (0.55, 0.55, 0.55)
        accent = (0.23, 0.49, 0.85)
        danger = (0.83, 0.21, 0.21)
        warn = (0.77, 0.48, 0.05)
        success = (0.10, 0.58, 0.32)

        sev_colors = {"high": danger, "medium": warn, "low": success}
        sev_labels = {"high": "ALTO", "medium": "MEDIO", "low": "BAJO"}

        x0, x1 = 50, 545
        y = 52

        def text(content, size=11, color=black, bold=False, x=x0):
            nonlocal y
            font = "helv" if not bold else "hebo"
            page.insert_text((x, y), content, fontsize=size, color=color, fontname=font)
            y += size * 1.55

        def hline(thickness=0.5, color=(0.82, 0.82, 0.82)):
            nonlocal y
            page.draw_line((x0, y), (x1, y), color=color, width=thickness)
            y += 8

        def spacer(h=6):
            nonlocal y
            y += h

        def maybe_new_page():
            nonlocal page, y
            if y > 780:
                page = doc.new_page(width=595, height=842)
                y = 52

        # ── Header ──────────────────────────────────────────────
        text("DIAGNÓSTICO LEGAL — ISSBC", size=9, color=light)
        now = datetime.now()
        text(now.strftime("%d/%m/%Y  %H:%M"), size=9, color=light, x=420)
        y -= 9 * 1.55  # align date on same line as header
        y += 9 * 1.55
        spacer(4)
        hline(thickness=1.0, color=accent)
        spacer(6)

        # ── Diagnosis title ──────────────────────────────────────
        title = diagnosis.get("diagnosis", "Sin título")
        text(title, size=17, bold=True)
        spacer(2)

        sev = diagnosis.get("severity", "medium")
        sev_color = sev_colors.get(sev, gray)
        text(f"Gravedad: {sev_labels.get(sev, '—')}", size=10, color=sev_color)
        spacer(4)
        hline()

        # ── Summary ──────────────────────────────────────────────
        text("RESUMEN", size=9, color=light, bold=True)
        spacer(2)
        summary = diagnosis.get("summary", "")
        for line in _wrap(summary, 90):
            maybe_new_page()
            text(line, size=11, color=black)
        spacer(10)

        # ── Justification sections ────────────────────────────────
        justification = diagnosis.get("justification", [])
        if justification:
            hline()
            text("JUSTIFICACIÓN", size=9, color=light, bold=True)
            spacer(4)
            for section in justification:
                maybe_new_page()
                heading = section.get("heading", "")
                body = section.get("body", "")
                text(heading, size=11, bold=True, color=black)
                spacer(1)
                for line in _wrap(body, 88):
                    maybe_new_page()
                    text("  " + line, size=10.5, color=gray)
                spacer(8)

        # ── Hypotheses ────────────────────────────────────────────
        if hypotheses:
            maybe_new_page()
            hline()
            text("HIPÓTESIS", size=9, color=light, bold=True)
            spacer(4)
            for h in sorted(hypotheses, key=lambda x: x.get("score", 0), reverse=True):
                maybe_new_page()
                pct = int(h.get("score", 0) * 100)
                hsev = h.get("severity", "medium")
                hc = sev_colors.get(hsev, gray)
                text(f"● {h.get('name', '')}  —  {pct}%", size=11, color=hc)
                detail = h.get("detail", "")
                if detail:
                    for line in _wrap(detail, 86):
                        maybe_new_page()
                        text("  " + line, size=10, color=gray)
                spacer(6)

        # ── Citations ─────────────────────────────────────────────
        citations = diagnosis.get("citations", [])
        if citations:
            maybe_new_page()
            hline()
            text("CITAS", size=9, color=light, bold=True)
            spacer(4)
            for ci in citations:
                maybe_new_page()
                doc_name = ci.get("doc", ci.get("label", ci.get("name", "")))
                page_num = ci.get("page")
                note = ci.get("note", ci.get("quote", ""))
                ref = f"[PDF] {doc_name}"
                if page_num:
                    ref += f"  -  p. {page_num}"
                text(ref, size=10.5, color=accent)
                if note:
                    for line in _wrap(note, 85):
                        maybe_new_page()
                        text("  " + line, size=10, color=gray)
                spacer(6)

        # ── Footer ────────────────────────────────────────────────
        page.draw_line((x0, 820), (x1, 820), color=(0.82, 0.82, 0.82), width=0.5)
        page.insert_text((x0, 832), "Generado por ISSBC · Diagnóstico Legal Asistido por IA",
                         fontsize=8, color=light, fontname="helv")

        doc.save(path)
        doc.close()

    def export_docx(self, path: str, diagnosis: dict, hypotheses: list) -> None:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        sev_labels = {"high": "ALTO", "medium": "MEDIO", "low": "BAJO"}
        sev_colors = {
            "high": RGBColor(0xD3, 0x35, 0x35),
            "medium": RGBColor(0xC5, 0x7A, 0x0D),
            "low": RGBColor(0x1A, 0x94, 0x52),
        }

        doc = Document()

        # narrow margins
        for section in doc.sections:
            section.top_margin = section.bottom_margin = Pt(48)
            section.left_margin = section.right_margin = Pt(56)

        # title
        title_par = doc.add_heading(diagnosis.get("diagnosis", "Sin título"), level=1)
        title_par.runs[0].font.size = Pt(18)

        # severity
        sev = diagnosis.get("severity", "medium")
        sev_par = doc.add_paragraph()
        run = sev_par.add_run(f"Gravedad: {sev_labels.get(sev, '—')}")
        run.font.color.rgb = sev_colors.get(sev, RGBColor(0x55, 0x55, 0x55))
        run.font.bold = True
        run.font.size = Pt(10)

        doc.add_paragraph()

        # summary
        doc.add_heading("Resumen", level=2)
        doc.add_paragraph(diagnosis.get("summary", ""))

        # justification
        justification = diagnosis.get("justification", [])
        if justification:
            doc.add_heading("Justificación", level=2)
            for section in justification:
                h = doc.add_paragraph()
                r = h.add_run(section.get("heading", ""))
                r.bold = True
                doc.add_paragraph(section.get("body", ""))

        # hypotheses
        if hypotheses:
            doc.add_heading("Hipótesis", level=2)
            for h in sorted(hypotheses, key=lambda x: x.get("score", 0), reverse=True):
                pct = int(h.get("score", 0) * 100)
                hsev = h.get("severity", "medium")
                p = doc.add_paragraph(style="List Bullet")
                name_run = p.add_run(f"{h.get('name', '')}  —  {pct}%")
                name_run.bold = True
                name_run.font.color.rgb = sev_colors.get(hsev, RGBColor(0x55, 0x55, 0x55))
                if h.get("detail"):
                    detail_p = doc.add_paragraph(h["detail"])
                    detail_p.paragraph_format.left_indent = Pt(18)

        # footer note
        doc.add_paragraph()
        foot = doc.add_paragraph()
        foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = foot.add_run(f"Generado por ISSBC · {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        doc.save(path)

    def export_csv(self, path: str, chat_history: list) -> None:
        with open(path, "w", newline="", encoding="utf-8-sig") as f:
            writer = csv.writer(f)
            writer.writerow(["timestamp", "role", "content"])
            for msg in chat_history:
                writer.writerow([
                    msg.get("timestamp", ""),
                    msg.get("role", ""),
                    msg.get("content", ""),
                ])

    def share_email(self, diagnosis: dict, hypotheses: list) -> None:
        sev_labels = {"high": "ALTO", "medium": "MEDIO", "low": "BAJO"}
        title = diagnosis.get("diagnosis", "Diagnóstico legal")
        sev = sev_labels.get(diagnosis.get("severity", "medium"), "")
        summary = diagnosis.get("summary", "")

        lines = [f"DIAGNÓSTICO: {title}", f"Gravedad: {sev}", "", summary, "", "HIPÓTESIS:"]
        for h in sorted(hypotheses, key=lambda x: x.get("score", 0), reverse=True):
            pct = int(h.get("score", 0) * 100)
            lines.append(f"- {h.get('name', '')} ({pct}%)")
            if h.get("detail"):
                lines.append(f"  {h['detail']}")
        lines += ["", "---", "Generado por ISSBC · Diagnóstico Legal Asistido por IA"]

        body = urllib.parse.quote("\n".join(lines))
        subject = urllib.parse.quote(f"Diagnóstico legal: {title}")
        webbrowser.open(f"mailto:?subject={subject}&body={body}")
